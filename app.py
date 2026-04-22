# Astra Resume Engine — Personalised for Akhila Reddy (v4.0)
# Aggressive ATS Optimisation Edition — "Get the Call, Handle the Interview"
# Target Roles: Senior Java Full Stack Developer, Senior Software Engineer,
#               Java Backend Engineer, Full Stack Engineer, Principal Engineer
import streamlit as st
import json
import re
import io
import ast
import datetime
from typing import List
from pydantic import BaseModel, Field
from google import genai
from google.genai import types
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
from reportlab.lib.units import inch
from xml.sax.saxutils import escape
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# API KEYS — loaded from Streamlit Secrets (never exposed to the user)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
try:
    google_key = st.secrets["GOOGLE_API_KEY"]
except Exception:
    google_key = ""
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# MODELS
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
GENERATION_MODEL = "gemini-3-flash-preview"
SCORING_MODEL = "gemini-3.1-flash-lite-preview"
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 1. CONFIGURATION
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
PAGE_TITLE = "Astra — Akhila Reddy"
AKHILA_BASE_RESUME = """Akhila Reddy
Senior Java Full Stack Developer
Phone: 9016094421 | Email: akhilareddy0213@gmail.com | LinkedIn: linkedin.com/in/akhila-reddy-835847184

PROFESSIONAL SUMMARY
Senior Java Full Stack Developer with 10+ years of experience designing, developing, and deploying enterprise applications using Java and J2EE technologies across Banking, Healthcare, Retail, Government, and Telecom domains. Strong hands-on expertise across full SDLC including requirement analysis, design, coding, testing, implementation, and maintenance. Proven track record building scalable microservices on AWS and Azure, modernizing legacy monoliths to Spring Boot + Angular/React SPAs, and integrating LLM/GenAI services (OpenAI, Claude) into enterprise workflows.

TECHNICAL SKILLS
Programming Languages: Java 8/11/17, C, C++, PL/SQL, TypeScript, JavaScript (ES6), Scala (contributed)
Java/J2EE & Frameworks: Spring Boot, Spring MVC, Spring Security, Spring Data, Spring Batch, Spring Cloud, Spring AOP, Spring IOC, Hibernate, JPA, MyBatis, iBatis, Spring JDBC, EJB, Servlets, JSP, JSTL, JSF, Struts 2
JavaScript Frameworks: Angular 8/10/14/16/17, AngularJS, React JS, Redux, React Router, Flux, Node.js, Express.js, jQuery
Web Technologies: HTML5, CSS3, AJAX, Bootstrap, XML, JSON, Webpack, Babel, NPM
Web Services/APIs: REST, SOAP, GraphQL, JAX-RS, JAX-WS, JAX-RPC, Apache CXF, Jersey, Apache Axis, HATEOAS, Swagger, Apigee, OAuth2, JWT
Messaging: Apache Kafka, RabbitMQ, IBM MQ
Databases: Oracle, SQL Server, MySQL, PostgreSQL, DB2, Sybase, MongoDB, Cassandra, DynamoDB, Redis
Cloud Platforms: AWS (EC2, S3, Lambda, RDS, DynamoDB, SQS, SNS, EKS, CloudWatch, Route53, VPC, IAM, API Gateway, Glue, Kinesis, CloudFormation, Elastic Beanstalk), Azure (Functions, Event Grid, Service Bus, Key Vault, App Services, Azure DevOps), GCP (BigQuery, GKE), Pivotal Cloud Foundry (PCF)
DevOps & CI/CD: Jenkins, Bamboo, AWS CodePipeline, Azure DevOps, GitHub Actions, Docker, Kubernetes, Terraform (IaC), Ansible, Helm
Testing: JUnit, JUnit5, Mockito, PowerMock, EasyMock, Spock, Selenium WebDriver, Cucumber, JMeter, Jasmine, Karma, Mocha, Chai, Jest, Postman, SOAPUI, Swagger UI
Monitoring & Observability: Splunk, Grafana, Kibana, New Relic, DataDog, ELK Stack, CloudWatch, Spring Boot Actuator
Build Tools: Maven, Gradle, ANT
Version Control: Git, GitHub, GitLab, Bitbucket, SVN
App Servers: IBM WebSphere, Oracle WebLogic, JBoss, Apache Tomcat
Other: Apache Camel, Mule ESB, Spring Integration, Drools Rule Engine, Camunda BPM, Jakarta XML Binding (JAXB), AI/GenAI (OpenAI API, Claude API, GitHub Copilot), Netflix OSS (Eureka, Hystrix), Shell Scripting (Bash, KSH), Agile/Scrum, JIRA
Methodologies: Agile, Scrum, Waterfall, OOAD, SOLID, Microservices, Event-Driven Architecture, Micro Frontends

PROFESSIONAL EXPERIENCE

Senior Java Full-Stack Developer | Truist Bank | Charlotte, NC | April 2025 - Present
- Designed and developed enterprise banking applications using Java 17, Spring Boot, Spring MVC, Hibernate, and RESTful Web Services across the full SDLC lifecycle.
- Built Angular 16/17 Single Page Applications with TypeScript, HTML5, CSS3, implementing modular architecture and reusable UI components using Angular Material, PrimeNG, and Kendo UI.
- Decomposed monolithic banking applications into Spring Boot microservices, improving scalability, maintainability, and deployment flexibility.
- Implemented secure API communication using Spring Security, OAuth2, and JWT authentication with role-based access control.
- Configured CI/CD pipelines using Jenkins, Maven, Git, Azure DevOps, and GitHub Actions for automated builds, testing, and deployment.
- Containerized applications using Docker and deployed services on Kubernetes for scalable microservices deployment.
- Designed and deployed secure REST APIs and Spring Boot microservices on Microsoft Azure, integrating with Apigee for API proxy management, traffic routing, and enterprise-grade exposure.
- Built and managed Apigee API proxies and policies for OAuth2, JWT validation, API key verification, rate limiting, and request/response transformation.
- Designed and optimized Oracle database schemas and PL/SQL procedures including complex queries, stored procedures, and triggers.
- Designed GraphQL schemas, resolvers, and query optimization strategies for high-performance APIs.
- Integrated LLM-based AI services (OpenAI and Claude APIs) into banking applications for intelligent automation and contextual decision-making.
- Built AI-assisted features for code generation, testing, and debugging optimization using GitHub Copilot and LLM-based assistants.
- Migrated legacy banking services to Azure-based microservices exposed through Apigee Hybrid, improving scalability and developer onboarding.
- Implemented thread-safe programming using Java multithreading, Executor Services, and synchronization mechanisms.
- Led backend API design discussions and mentored junior developers on microservices and API best practices.
- Provided L2/L3 production support, triaging issues, performing root cause analysis, and resolving critical defects.
Environment: Java 17, Spring Boot, Spring Security, Angular 17, REST APIs, GraphQL, Oracle DB, PostgreSQL, Kafka, Docker, Kubernetes, Jenkins, Git, Swagger, Splunk, Apigee, Azure, AWS.

Senior Java Full Stack Developer | Centene Corporation | Saint-Louis, MO | January 2024 - April 2025
- Developed responsive web applications using HTML5, CSS3, jQuery, Angular 14, JavaScript, Node.js, and JSON for healthcare platforms.
- Implemented strong typing and interface-based architecture in TypeScript to improve code quality, maintainability, and scalability.
- Contributed to services implemented in Scala to enhance performance and adopt functional programming practices.
- Migrated legacy Java applications to Java 17 with enhanced performance and security, implementing Java 17 records and sealed classes.
- Designed Infrastructure as Code (IaC) using Terraform to provision AWS resources including EC2, Lambda, API Gateway, DynamoDB, RDS, IAM roles, and VPC.
- Developed reusable Terraform modules for multi-environment (dev, QA, prod) deployments integrated with CI/CD pipelines.
- Automated AWS EC2 provisioning using Ansible playbooks and YAML scripts.
- Designed event-driven microservices using Spring Boot and Apache Kafka deployed on AWS.
- Implemented ETL pipelines using AWS Glue to extract, transform, and load large datasets into cloud data stores.
- Integrated microservices with AWS Glue pipelines for downstream data processing and loading into cloud data stores.
- Developed HATEOAS-based RESTful web services using Spring Boot frameworks.
- Implemented secure REST APIs using Spring Security and OAuth2 authentication mechanisms.
- Designed observability dashboards using Grafana, Splunk, and CloudWatch, defining SLIs/SLOs and alert thresholds.
- Built unit and integration tests using JUnit5, Mockito, and Spring Boot Test, improving microservice code coverage.
- Developed micro frontend architecture using Angular and module federation, enabling independent deployment and scalability of UI components.
- Deployed applications on Azure cloud (Azure Functions, Service Bus, App Services) with CI/CD pipelines and cloud-native integrations.
- Leveraged AI-assisted development tools (GitHub Copilot, LLM-based code assistants) for automated code generation, refactoring, and test case generation.
- Implemented business rules using Drools rule engine and rule flows for healthcare claims processing.
- Performed performance and load testing using Apache JMeter.
- Configured Karma test coverage reports to monitor Angular code quality.
- Worked on Docker-based Jenkins setup with Kubernetes-managed containers for CI/CD orchestration.
Environment: Java 11/17, Spring Boot, Spring MVC, Spring Security, Hibernate ORM, Angular 14, TypeScript, RESTful APIs, Apache Kafka, Drools Rule Engine, OAuth2, Oracle DB, MongoDB, Terraform, Jenkins, Docker, Kubernetes, AWS (EC2, Lambda, S3, Glue), JUnit5, Mockito, Grafana, Splunk, Apigee.

Java Full Stack Developer | Walmart | Sunnyvale, CA | January 2022 - January 2024
- Developed enterprise retail applications using Java, JSP, JDBC, HTML, XML, JavaScript, and JUnit.
- Extensively wrote and optimized SQL and PL/SQL queries in Sybase and SQL Server databases.
- Designed modular and reusable UI components using TypeScript and Angular 8.
- Built microservices using Docker containers, AWS Lambda, and Elastic Beanstalk.
- Developed microservices-based applications using Spring Boot, Spring Cloud, Netflix OSS (Eureka, Hystrix), RabbitMQ, Kafka, and RESTful services.
- Implemented messaging solutions using RabbitMQ and Kafka for asynchronous communication.
- Worked with GCP services including BigQuery and GKE for large-scale data processing and containerized deployments.
- Designed multi-cloud architecture integrating AWS microservices with GCP-based analytics platforms.
- Implemented Apache Kafka producers and consumers integrated with Zookeeper clusters.
- Optimized MongoDB CRUD operations for better performance and scalability.
- Developed automation scripts using Korn Shell (KSH) and Bash on UNIX/Linux environments for deployment, log monitoring, and batch job scheduling.
- Designed automation for microservices infrastructure, single sign-on, MFA security, and access management.
- Applied SOLID principles and OOPS concepts to build scalable and maintainable systems.
- Utilized JAXB (Jakarta XML Binding) for marshalling and unmarshalling XML data in REST/SOAP services.
- Migrated applications and services to AWS cloud environments.
Environment: Java 11, Spring Boot, Spring MVC, Spring Security, Hibernate, Microservices Architecture, Angular 8, TypeScript, RESTful APIs, Kafka, RabbitMQ, Redis, Oracle DB, SQL Server, Maven, Jenkins, Docker, AWS (EC2, Lambda, Elastic Beanstalk), GCP (BigQuery, GKE), Git, JUnit, Karma, Jasmine, Linux/Unix Shell Scripting.

Java Full Stack Developer | State of New Mexico Supreme Court | Santa Fe, NM | February 2020 - January 2022
- Designed overall application layouts and prototypes using React JS, Redux, jQuery UI, HTML5, and CSS/Less.
- Built responsive UI components using React JS and Bootstrap, dynamically rendering data through REST APIs and Virtual DOM.
- Developed reusable React forms with validation logic using Redux architecture and modular components.
- Implemented Node.js as a server-side proxy for event-driven and non-blocking I/O applications.
- Utilized Webpack module bundler and Babel compiler for React JS application builds.
- Developed unit tests for TypeScript and JavaScript components using Jest, Karma, and Jasmine frameworks.
- Implemented Spring Boot-based CRUD backend services to support React front-end applications.
- Built and deployed microservices using Docker and Kubernetes orchestration.
- Implemented serverless architectures using AWS Lambda, Kinesis, and CloudFormation.
- Deployed Spring Boot applications to Pivotal Cloud Foundry (PCF) environments.
- Integrated Java Spring Boot microservices with Scala-based backend services using Akka HTTP APIs.
- Designed and implemented workflow-driven microservices using Camunda BPM, enabling orchestration of business processes.
- Integrated Camunda with Spring Boot microservices to manage long-running workflows and event-driven process execution.
- Performed integration testing using Selenium WebDriver and Cucumber.
- Automated front-end build processes using Gulp and JavaScript build tools.
Environment: Java 8, Spring Boot, Spring Security, Spring MVC, React JS, Redux, JavaScript (ES6), TypeScript, HTML5, CSS3, Bootstrap, RESTful APIs, Node.js, Express.js, Webpack, Babel, Oracle DB, MongoDB, Redis, Maven, Jenkins, Docker, Kubernetes, AWS (EC2, Lambda), Camunda BPM.

Senior Java Developer | Cisco | San Jose, CA | September 2018 - February 2020
- Developed dynamic web pages using JSP, JSF, and JSTL tag libraries for telecom/networking applications.
- Implemented Apache Kafka for real-time streaming and application metrics collection.
- Built Single Page Applications using Angular with two-way data binding and modular architecture.
- Designed highly scalable serverless microservices using AWS API Gateway, Lambda, and DynamoDB.
- Worked with Cassandra databases using CQL to design and execute efficient read and write operations, optimizing data models for high availability in distributed environments.
- Designed, developed, and consumed RESTful web services for distributed applications.
- Developed SOAP and REST web services using JAX-WS and Java XML APIs.
- Wrote complex PL/SQL queries and procedures using Oracle databases.
- Implemented middle-tier business logic using Spring Framework, Spring MVC, and Spring Boot.
- Developed reusable components using Spring AOP and IOC dependency injection.
- Worked on telecom domain applications supporting network services, provisioning, and real-time data processing.
- Developed backend services for telecom workflows (OSS/BSS systems) integrating APIs for service activation and billing systems.
Environment: Java 11, Spring Boot, Spring Cloud, AngularJS, Kafka, Cassandra, AWS Lambda, DynamoDB, API Gateway, Maven, Git, Docker.

Java Developer | IBM | Hyderabad, India | September 2013 - September 2016
- Developed application modules using Spring MVC, Spring Annotations, Spring Beans, Dependency Injection, and Spring AOP with Hibernate ORM.
- Implemented Spring Bean Factory to create proxy objects using AOP framework.
- Used Hibernate as ORM mapping tool and worked with Hibernate Query Language (HQL).
- Produced and consumed web services data in JSON and XML formats.
- Developed interactive web pages using HTML, CSS, JSP, JavaScript, jQuery, and AJAX.
- Designed and developed web tier components and RESTful APIs using Spring MVC.
- Developed functions and stored procedures using SQL and PL/SQL for Oracle database.
- Used Git for source control and Jenkins for continuous integration.
- Worked on JIRA for user requirements and bug tracking.
- Practiced Agile methodology for sprint tracking and delivery.
Environment: Java/J2EE, Spring MVC, Spring Integration, Hibernate 3.0, SOAP, RESTful, HTML, CSS3, JavaScript, jQuery, Eclipse, XML, AJAX, JBoss EAP 6.1, JUnit 4, Log4j, Git, JIRA, Oracle SQL/PL-SQL.

EDUCATION
Bachelor of Technology (B.Tech), Computer Science and Technology
Vaagdevi College of Engineering, Telangana, India | June 2009 - June 2013"""
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# ASTRA PROMPT v5.1 — The 98%+ ATS Mirror Engine (Company-Agnostic)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
ASTRA_PROMPT = """
Role: You are Astra, an elite ATS Bypass Engine. Your absolute directive is to generate a resume that scores 98-100% against the provided Job Description (JD) on any standard ATS scanner (Taleo, Workday, Greenhouse).

Candidate Base Profile: Akhila Reddy — Senior Java Full Stack Developer (10+ years). 
Roles: Truist (2025-Present), Centene (2024-2025), Walmart (2022-2024), NM Supreme Court (2020-2022), Cisco (2018-2020), IBM (2013-2016).

=== THE 98%+ ATS ALIGNMENT ALGORITHM (CRITICAL) ===

To achieve a 98%+ match, you must execute the following 4 directives without exception:

1. TOTAL ECOSYSTEM DOMINATION (NO DILUTION):
   - If the JD requires a specific framework (e.g., React), you MUST completely erase all mentions of competing frameworks (e.g., Angular, Vue, JSF, JSP) from her history. Rewrite those past projects to use the JD's requested framework. 
   - If the JD asks for a specific cloud (e.g., AWS), convert all GCP/Azure deployments in her past to AWS. 
   - Do not let competing technologies dilute the keyword density. 

2. VERBATIM PHRASE MIRRORING:
   - Do not just extract single keywords. Extract exact multi-word phrases from the JD (e.g., "AI ML infrastructure", "high-volume telemetry data", "cross-browser compatibility").
   - Inject these EXACT phrases verbatim into the Summary, Skills, and Experience bullet points. Do not paraphrase. ATS scanners look for exact string matches.

3. KEYWORD DENSITY DISTRIBUTION:
   - A keyword must not just appear in the "Skills" section. 
   - The top 5 core technologies from the JD MUST appear in the Summary, in the Skills section, AND be distributed across at least 3 different past jobs in the Experience section. Prove she has used these tools for years, not just recently.

4. SENIORITY & ARCHITECTURAL COMPLEXITY:
   - Even if the JD is for a pure Frontend or pure Backend role, maintain her 10+ years of seniority. 
   - Never dumb down her bullets. Use aggressive, architectural verbs: Architected, Orchestrated, Engineered, Spearheaded, Shipped.
   - Attach the JD's required skills to high-impact outcomes (e.g., "Reduced latency by 40% by implementing [JD Tool]").

=== WRITING CONSTRAINTS (SOUND HUMAN & PROFESSIONAL) ===
- BANNED WORDS: "testament to", "underscores", "pivotal", "realm", "landscape", "serves as", "showcasing", "leveraging", "seamless", "robust", "innovative".
- COMPANY-AGNOSTIC RULE: NEVER write the target company's name in the Summary or anywhere in the resume text. The summary must focus purely on her skills and architectural achievements.
- Bullet points must be dense with technical architecture, avoiding corporate fluff.

=== OUTPUT STRUCTURE ===
1. SUMMARY: 3-4 sentences. Sentence 1 MUST include her years of experience and the EXACT Job Title from the JD. Must contain 4-5 verbatim phrases from the JD. Do NOT mention the target company name here.
2. SKILLS: Categorized logically. The JD's exact requirements must be the very first items listed in each category.
3. EXPERIENCE: ALL 6 roles must be included. Rewrite every bullet point to act as a mirror to the JD's responsibilities and requirements.
4. TARGET COMPANY: Extract the company name from the JD for metadata purposes only.
5. MISSING KEYWORDS: This MUST be empty. You are required to find a logical way to weave 100% of the JD's requirements into her 10-year history. Do not leave any skill behind.
"""
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# COVER LETTER PROMPT — Aggressive Alignment Edition
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
COVER_LETTER_PROMPT = """
Role: You are Akhila Reddy, a Senior/Principal Java Full Stack Developer with 10+ years of experience, writing a direct email to a Hiring Manager.
Goal: Sound 100% human, supremely confident, and secure an interview by proving you have already solved the exact problems listed in their Job Description (JD).

=== AGGRESSIVE ALIGNMENT RULES ===
1. ASSUME TOTAL COMPETENCE: You have 10+ years of experience. You have handled whatever architecture, stack, or domain the JD is asking for. 
2. DOMAIN DOMINANCE: If the JD is in a domain you haven't explicitly worked in (e.g., Logistics, Automotive, Entertainment), aggressively adapt your past enterprise experience (Truist, Centene, Walmart) to sound like the perfect architectural precursor to their current problems. Never apologize for a lack of domain experience; project total authority.
3. TECH STACK INJECTION: Identify the top 3 or 4 most critical, difficult technologies in the JD (e.g., Kafka Streams, EKS, GraphQL, Terraform) and seamlessly name-drop them in your war story.

=== ANTI-AI-WRITING RULES (CRITICAL) ===
BANNED PHRASES — never use any of these:
"I am writing to express my interest", "I am excited to apply", "Please find my resume attached",
"testament to", "underscores", "pivotal", "realm", "tapestry", "I believe I am a perfect fit",
"passionate about", "driven by a desire", "committed to excellence", "at the forefront of",
"showcasing", "highlighting", "demonstrating", "serves as", "stands as", "fostering",
"leveraging", "harnessing", "utilizing", "seamless", "innovative", "groundbreaking"

WRITING STYLE:
- Short, punchy sentences mixed with medium ones. Vary the rhythm.
- Use plain, aggressive verbs: built, architected, shipped, migrated, scaled.
- No em dashes. Use commas or periods.
- Sound like a busy, senior engineer talking to another senior engineer. Cut the corporate fluff.

=== STRUCTURE ===
1. "Dear Hiring Team," (or use the company name, e.g., "Dear [Company] Engineering Team,")
2. THE HOOK: Open immediately with an observation about the specific technical challenge or business goal mentioned in the JD. 
   - Bad: "I am applying for the Senior Java role..."
   - Good: "Scaling a monolithic data pipeline into event-driven Spring Boot microservices while maintaining zero downtime is a massive undertaking."
3. THE WAR STORY (Dynamic): Synthesize a past project from your history (drawing on your Truist, Centene, or Walmart stacks) that perfectly mirrors their current challenge. Use the exact tech keywords from their JD. Tell them how you architected it, what tools you used, and how it solved the problem.
4. THE PITCH: One crisp sentence stating that with your 10+ years of enterprise Java/Full Stack experience, you can step in and drive this exact initiative for them.
5. CLOSING: "I'd love to jump on a call to discuss the architecture. Thank you."
6. SIGN OFF: "Akhila Reddy"

Return ONLY the letter body. No markdown. No bold. No headers.
"""
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# ATS SCORING PROMPT
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
ATS_SCORING_PROMPT = """You are a strict ATS (Applicant Tracking System) scanner.
Compare the RESUME JSON against the JOB DESCRIPTION.
Scoring criteria (0-100):
- Keyword match density (45%): What percentage of hard skills/tools/frameworks in the JD appear in the resume?
- Experience relevance (25%): Do the bullet points describe work that solves the JD's problems?
- Seniority alignment (15%): Does the experience level match what the JD asks for (10+ years senior/lead/principal)?
- Domain fit (15%): Is the candidate's industry background relevant or plausibly transferable?
Target: 98%+ for a well-tailored resume.
Output ONLY valid JSON with no markdown, no backticks, no explanation:
{"score": <int 0-100>, "reasoning": "<1 sentence>", "missing_keywords": "<comma-separated list of JD keywords NOT found in resume>"}
"""
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 2. PYDANTIC SCHEMAS
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
class ExperienceItem(BaseModel):
    role_title: str = Field(description="The job title exactly as it should appear")
    company: str = Field(description="The company name")
    dates: str = Field(description="Employment dates")
    location: str = Field(description="City, State/Country")
    responsibilities: List[str] = Field(description="List of 6-10 bullet points reframed for the JD with maximum keyword density")
    achievements: List[str] = Field(description="List of 0-3 achievements. ONLY include if grounded in the base resume. Never invent metrics.")
class EducationItem(BaseModel):
    degree: str = Field(description="Full degree name")
    college: str = Field(description="University name")
class SkillCategory(BaseModel):
    category: str = Field(description="Skill category name (e.g., 'Backend & Frameworks')")
    technologies: str = Field(description="Comma-separated tools. JD-mentioned tools listed FIRST.")
class ResumeSchema(BaseModel):
    candidate_name: str = Field(description="Always: Akhila Reddy")
    candidate_title: str = Field(description="Professional title tailored to match the JD's exact role title (e.g., 'Senior Java Full Stack Developer', 'Lead Software Engineer', 'Principal Backend Engineer')")
    contact_info: str = Field(description="Always: 9016094421 | akhilareddy0213@gmail.com | linkedin.com/in/akhila-reddy-835847184")
    summary: str = Field(description="3-4 sentence professional summary. Must mention 10+ years, target company name, and match JD language.")
    skills: List[SkillCategory] = Field(description="6-8 dense skill categories. Every JD keyword must appear here.")
    experience: List[ExperienceItem] = Field(description="ALL 6 roles: Truist, Centene, Walmart, NM Supreme Court, Cisco, IBM. Never drop any.")
    education: List[EducationItem] = Field(description="B.Tech CS from Vaagdevi College of Engineering")
    target_company: str = Field(description="Company name extracted from JD")
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SCHEMA CLEANER (removes fields Gemini API rejects)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def get_clean_schema(pydantic_cls):
    schema = pydantic_cls.model_json_schema()
    def _clean(d):
        if isinstance(d, dict):
            for key in ["additionalProperties", "title"]:
                d.pop(key, None)
            for v in d.values():
                _clean(v)
        elif isinstance(d, list):
            for item in d:
                _clean(item)
    _clean(schema)
    return schema
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 3. DATA NORMALIZER
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def clean_skill_string(skill_str):
    if not isinstance(skill_str, str):
        return str(skill_str)
    if skill_str.strip().startswith("["):
        try:
            list_match = re.search(r"\[(.*?)\]", skill_str)
            if list_match:
                actual_list = ast.literal_eval(list_match.group(0))
                extra_part = skill_str[list_match.end():].strip().lstrip(",").strip()
                clean_str = ", ".join([str(s) for s in actual_list])
                if extra_part:
                    clean_str += f", {extra_part}"
                return clean_str
        except Exception:
            pass
    return skill_str
def normalize_schema(data):
    if not isinstance(data, dict):
        return {"summary": str(data), "skills": {}, "experience": []}
    normalized = {}
    # Contact/Name
    normalized['candidate_name'] = data.get('candidate_name', 'Akhila Reddy')
    normalized['candidate_title'] = data.get('candidate_title', 'Senior Java Full Stack Developer')
    raw_contact = data.get('contact_info', '9016094421 | akhilareddy0213@gmail.com | linkedin.com/in/akhila-reddy-835847184')
    normalized['contact_info'] = str(raw_contact) if not isinstance(raw_contact, dict) else ' | '.join(str(v) for v in raw_contact.values() if v)
    # Summary
    normalized['summary'] = data.get('summary', '')
    # Skills → always dict
    raw_skills = data.get('skills', {})
    normalized['skills'] = {}
    if isinstance(raw_skills, dict):
        for k, v in raw_skills.items():
            normalized['skills'][k] = clean_skill_string(str(v))
    elif isinstance(raw_skills, list):
        for item in raw_skills:
            if isinstance(item, dict):
                cat = item.get('category', '')
                tech = item.get('technologies', '')
                if cat and tech:
                    normalized['skills'][cat] = clean_skill_string(str(tech))
            else:
                normalized['skills'] = {"General Skills": ", ".join([str(s) for s in raw_skills])}
                break
    # Experience
    raw_exp = data.get('experience', [])
    norm_exp = []
    if isinstance(raw_exp, list):
        for role in raw_exp:
            if isinstance(role, dict):
                norm_exp.append({
                    'role_title': role.get('role_title', ''),
                    'company': role.get('company', ''),
                    'dates': role.get('dates', ''),
                    'location': role.get('location', ''),
                    'responsibilities': role.get('responsibilities', []),
                    'achievements': role.get('achievements', []),
                })
    normalized['experience'] = norm_exp
    # Education
    raw_edu = data.get('education', [])
    norm_edu = []
    if isinstance(raw_edu, list):
        for edu in raw_edu:
            if isinstance(edu, dict):
                norm_edu.append({
                    'degree': edu.get('degree', ''),
                    'college': edu.get('college', ''),
                })
            elif isinstance(edu, str):
                norm_edu.append({'degree': edu, 'college': ''})
    elif isinstance(raw_edu, str):
        norm_edu.append({'degree': raw_edu, 'college': ''})
    if not norm_edu:
        norm_edu = [{'degree': 'Bachelor of Technology (B.Tech), Computer Science and Technology', 'college': 'Vaagdevi College of Engineering, TS, India'}]
    normalized['education'] = norm_edu
    normalized['target_company'] = data.get('target_company', 'Company')
    return normalized
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 4. ATS SCORING (Gemini 3.1 Flash-Lite)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def calculate_ats_score(resume_json, jd_text, api_key):
    if not api_key:
        return {"score": 0, "reasoning": "No API Key", "missing_keywords": ""}
    client = genai.Client(api_key=api_key)
    try:
        response = client.models.generate_content(
            model=SCORING_MODEL,
            contents=f"{ATS_SCORING_PROMPT}\n\nRESUME:\n{str(resume_json)[:4000]}\n\nJOB DESCRIPTION:\n{jd_text[:4000]}",
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
            )
        )
        content = response.text.strip()
        # Clean markdown fences if present
        if "```" in content:
            match = re.search(r"```(?:json)?(.*?)```", content, re.DOTALL)
            if match:
                content = match.group(1).strip()
        return json.loads(content)
    except Exception as e:
        return {"score": 0, "reasoning": f"Scoring Error: {str(e)}", "missing_keywords": ""}
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 5. JAVA FULL-STACK-SPECIFIC SKILL EXPANSION
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def expand_skills_dense(skills):
    """For any tool she's used, auto-add closely related variants to increase keyword density."""
    if not skills:
        return {}
    EXPANSIONS = {
        "Spring Boot": "Spring WebFlux, Spring Cloud Gateway, Spring Batch",
        "Kafka": "Kafka Streams, Kafka Connect, Confluent Platform",
        "Kubernetes": "Helm, K8s Deployments, kubectl, Kustomize",
        "Docker": "Docker Compose, Container Registry, OCI Images",
        "Jenkins": "Jenkinsfile, Pipeline as Code",
        "Terraform": "Terragrunt, Terraform Modules, Terraform Cloud",
        "AWS": "AWS CloudFormation, AWS IAM, AWS CDK",
        "Azure": "Azure Resource Manager (ARM), Azure Pipelines",
        "GCP": "GCP IAM, GCP Cloud Run",
        "Oracle": "Oracle SQL Developer, Oracle PL/SQL",
        "MongoDB": "MongoDB Atlas, Mongoose",
        "Angular": "Angular CLI, RxJS, NgRx",
        "React JS": "React Hooks, React Context API",
        "Node.js": "npm, Yarn, ExpressJS middleware",
        "Hibernate": "Hibernate Envers, Hibernate Search",
        "Maven": "Maven Central, Maven Wrapper",
        "Grafana": "Prometheus, Loki",
        "Splunk": "Splunk SPL, Splunk Dashboards",
        "Apigee": "Apigee Edge, Apigee Hybrid, API Proxies",
        "GraphQL": "Apollo, GraphQL Federation",
        "Selenium": "Selenium Grid, TestNG",
        "JUnit": "AssertJ, Hamcrest",
        "REST": "OpenAPI 3.0, RESTful Design",
        "SOAP": "WSDL, XSD Validation",
        "Camunda": "BPMN 2.0, DMN",
    }
    for cat, tools in skills.items():
        tools_str = str(tools)
        for trigger, additions in EXPANSIONS.items():
            if trigger in tools_str:
                for addition in additions.split(", "):
                    if addition not in tools_str:
                        tools_str += f", {addition}"
        skills[cat] = tools_str
    return skills
def to_text_block(val):
    if val is None:
        return ""
    if isinstance(val, list):
        return "\n".join([str(x) for x in val])
    return str(val)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 6. GENERATION (Gemini 3 Flash)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def analyze_and_generate(api_key, resume_text, jd_text):
    client = genai.Client(api_key=api_key)
    try:
        safe_schema = get_clean_schema(ResumeSchema)
        response = client.models.generate_content(
            model=GENERATION_MODEL,
            contents=f"{ASTRA_PROMPT}\n\nRESUME:\n{resume_text}\n\nJD:\n{jd_text}",
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=safe_schema,
            )
        )
        raw_data = json.loads(response.text)
        data = raw_data.model_dump() if hasattr(raw_data, 'model_dump') else raw_data
        # Transform skills list → dict if needed
        if 'skills' in data and isinstance(data['skills'], list):
            transformed = {}
            for item in data['skills']:
                cat = item.get('category') if isinstance(item, dict) else getattr(item, 'category', '')
                tech = item.get('technologies') if isinstance(item, dict) else getattr(item, 'technologies', '')
                if cat and tech:
                    transformed[cat] = tech
            data['skills'] = transformed
        data = normalize_schema(data)
        data['skills'] = expand_skills_dense(data.get('skills', {}))
        # ATS Score using the lite model
        judge = calculate_ats_score(data, jd_text, api_key)
        data['ats_score'] = judge.get('score', 0)
        data['ats_reason'] = judge.get('reasoning', '')
        data['missing_keywords'] = judge.get('missing_keywords', '')
        return data
    except Exception as e:
        return {"error": f"Generation Error: {str(e)}"}
def generate_cover_letter(api_key, resume_data, jd_text):
    client = genai.Client(api_key=api_key)
    try:
        response = client.models.generate_content(
            model=GENERATION_MODEL,
            contents=f"{COVER_LETTER_PROMPT}\n\nRESUME DATA:\n{str(resume_data)}\n\nJOB DESCRIPTION:\n{jd_text}",
        )
        return response.text
    except Exception as e:
        return f"Error generating cover letter: {str(e)}"
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 7. DOCX RENDERER
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def set_font(run, size, bold=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    try:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    except Exception:
        pass
def create_doc(data):
    doc = Document()
    s = doc.sections[0]
    s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(0.5)
    # Header
    for txt, sz, b in [
        (data.get('candidate_name', ''), 28, True),
        (data.get('candidate_title', ''), 14, True),
        (data.get('contact_info', ''), 12, True),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(to_text_block(txt))
        if sz == 28:
            run.font.all_caps = True
        set_font(run, sz, b)
    def add_sec(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(2)
        set_font(p.add_run(title), 12, True)
    def add_body(txt, bullet=False):
        style = 'List Bullet' if bullet else 'Normal'
        p = doc.add_paragraph(style=style)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(to_text_block(txt)), 12)
    # Professional Profile
    add_sec("Professional Profile")
    add_body(data.get('summary', ''))
    # Skills
    add_sec("Technical Skills")
    for k, v in data.get('skills', {}).items():
        p = doc.add_paragraph(style='List Bullet')
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(f"{k}: "), 12, True)
        set_font(p.add_run(to_text_block(v)), 12)
    # Experience
    add_sec("Professional Experience")
    for role in data.get('experience', []):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        line = f"{role.get('role_title')} | {role.get('company')} | {role.get('location')} | {role.get('dates')}"
        set_font(p.add_run(to_text_block(line)), 12, True)
        resps = role.get('responsibilities', [])
        if isinstance(resps, str):
            resps = resps.split('\n')
        for r in resps:
            if str(r).strip():
                add_body(r, bullet=True)
        achs = role.get('achievements', [])
        if isinstance(achs, str):
            achs = achs.split('\n')
        if achs and any(str(a).strip() for a in achs):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            set_font(p.add_run("Achievements:"), 12, True)
            for a in achs:
                if str(a).strip():
                    add_body(a, bullet=True)
    # Education
    add_sec("Education")
    for edu in data.get('education', []):
        text = f"{edu.get('degree', '')}, {edu.get('college', '')}"
        add_body(text, bullet=True)
    return doc
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 8. COVER LETTER DOCX
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def create_cover_letter_doc(cover_letter_text, data):
    doc = Document()
    s = doc.sections[0]
    s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(0.5)
    def add_line(text, bold=False, space_after=12, align=WD_PARAGRAPH_ALIGNMENT.LEFT):
        if not text:
            return
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(str(text))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = bold
    add_line(data.get('candidate_name', '').upper(), bold=True, space_after=0)
    contact_info = data.get('contact_info', '')
    if "|" in contact_info:
        for part in contact_info.split('|'):
            add_line(part.strip(), bold=False, space_after=0)
    else:
        add_line(contact_info, bold=False, space_after=0)
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    today_str = datetime.date.today().strftime("%B %d, %Y")
    add_line(today_str, space_after=12)
    for para in cover_letter_text.split('\n'):
        if para.strip():
            add_line(para.strip(), bold=False, space_after=12, align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    return doc
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 9. PDF RENDERER
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def create_pdf(data):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=letter,
        leftMargin=0.5 * inch, rightMargin=0.5 * inch,
        topMargin=0.5 * inch, bottomMargin=0.5 * inch,
    )
    styles = getSampleStyleSheet()
    sn = ParagraphStyle('N', parent=styles['Normal'], fontName='Times-Roman', fontSize=12, leading=14, alignment=TA_JUSTIFY, spaceAfter=0)
    sh_name = ParagraphStyle('HName', parent=styles['Normal'], fontName='Times-Bold', fontSize=28, leading=30, alignment=TA_CENTER, spaceAfter=0)
    sh_title = ParagraphStyle('HTitle', parent=styles['Normal'], fontName='Times-Bold', fontSize=14, leading=16, alignment=TA_CENTER, spaceAfter=0)
    sh_contact = ParagraphStyle('HContact', parent=styles['Normal'], fontName='Times-Bold', fontSize=12, leading=14, alignment=TA_CENTER, spaceAfter=6)
    s_sec = ParagraphStyle('Sec', parent=styles['Normal'], fontName='Times-Bold', fontSize=12, leading=14, alignment=TA_LEFT, spaceBefore=12, spaceAfter=2)
    def clean(txt):
        if txt is None:
            return ""
        txt = to_text_block(txt)
        return escape(txt).replace('\n', '<br/>')
    elements = []
    elements.append(Paragraph(clean(data.get('candidate_name', '')), sh_name))
    elements.append(Paragraph(clean(data.get('candidate_title', '')), sh_title))
    elements.append(Paragraph(clean(data.get('contact_info', '')), sh_contact))
    elements.append(Paragraph("Professional Profile", s_sec))
    elements.append(Paragraph(clean(data.get('summary', '')), sn))
    elements.append(Paragraph("Technical Skills", s_sec))
    skill_items = []
    for k, v in data.get('skills', {}).items():
        text = f"<b>{clean(k)}:</b> {clean(v)}"
        skill_items.append(ListItem(Paragraph(text, sn), leftIndent=0))
    if skill_items:
        elements.append(ListFlowable(skill_items, bulletType='bullet', start='\u2022', leftIndent=15))
    elements.append(Paragraph("Professional Experience", s_sec))
    for role in data.get('experience', []):
        line = f"{role.get('role_title')} | {role.get('company')} | {role.get('location')} | {role.get('dates')}"
        elements.append(Paragraph(f"<b>{clean(line)}</b>", sn))
        elements.append(Spacer(1, 2))
        role_bullets = []
        resps = role.get('responsibilities', [])
        if isinstance(resps, str):
            resps = resps.split('\n')
        for r in resps:
            if str(r).strip():
                role_bullets.append(ListItem(Paragraph(clean(r), sn), leftIndent=0))
        if role_bullets:
            elements.append(ListFlowable(role_bullets, bulletType='bullet', start='\u2022', leftIndent=15))
        achs = role.get('achievements', [])
        if isinstance(achs, str):
            achs = achs.split('\n')
        if achs and any(str(a).strip() for a in achs):
            elements.append(Paragraph("<b>Achievements:</b>", sn))
            ach_bullets = []
            for a in achs:
                if str(a).strip():
                    ach_bullets.append(ListItem(Paragraph(clean(a), sn), leftIndent=0))
            if ach_bullets:
                elements.append(ListFlowable(ach_bullets, bulletType='bullet', start='\u2022', leftIndent=25))
        elements.append(Spacer(1, 6))
    elements.append(Paragraph("Education", s_sec))
    edu_bullets = []
    for edu in data.get('education', []):
        text = f"{edu.get('degree', '')}, {edu.get('college', '')}"
        edu_bullets.append(ListItem(Paragraph(clean(text), sn), leftIndent=0))
    if edu_bullets:
        elements.append(ListFlowable(edu_bullets, bulletType='bullet', start='\u2022', leftIndent=15))
    doc.build(elements)
    buffer.seek(0)
    return buffer.getvalue()
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 10. STREAMLIT UI
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
st.set_page_config(page_title=PAGE_TITLE, layout="wide", page_icon="\U0001f680", initial_sidebar_state="expanded")
st.markdown("""
<style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    .block-container {padding-top: 1.5rem;}
    div.stButton > button:first-child {border-radius: 6px; font-weight: 600;}
    div[data-testid="stMetricValue"] {font-size: 1.8rem;}
</style>
""", unsafe_allow_html=True)
if 'data' not in st.session_state:
    st.session_state['data'] = None
if 'saved_base' not in st.session_state:
    st.session_state['saved_base'] = AKHILA_BASE_RESUME
if 'saved_jd' not in st.session_state:
    st.session_state['saved_jd'] = ""
if 'cover_letter' not in st.session_state:
    st.session_state['cover_letter'] = None
with st.sidebar:
    st.header("\u2699\ufe0f Configuration")
    if google_key:
        st.success("API key configured")
    else:
        st.error("API key missing — add GOOGLE_API_KEY to Streamlit secrets")
        google_key = st.text_input("Google API Key (fallback)", type="password")
    st.divider()
    st.markdown("**Target Roles:**")
    st.caption("Senior Java Full Stack \u2022 Lead Software Engineer \u2022 Principal Backend \u2022 Java Microservices")
    st.divider()
    st.markdown("**Models:**")
    st.caption(f"Resume: {GENERATION_MODEL}")
    st.caption(f"Scoring: {SCORING_MODEL}")
    st.divider()
    if st.button("\U0001f5d1\ufe0f Reset", use_container_width=True):
        st.session_state['data'] = None
        st.session_state['saved_base'] = AKHILA_BASE_RESUME
        st.session_state['saved_jd'] = ""
        st.session_state['cover_letter'] = None
        st.rerun()
    st.caption("Astra v4.0 | Personalised for Akhila")
if not st.session_state['data']:
    st.markdown(f"<h1 style='text-align: center;'>{PAGE_TITLE}</h1>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; color: #888;'>Paste a JD. Get a 98% tailored resume. Get the call.</p>", unsafe_allow_html=True)
    st.divider()
    c1, c2 = st.columns(2)
    with c1:
        st.subheader("\U0001f4cb Base Resume")
        st.caption("Pre-loaded. Edit only if needed.")
        base = st.text_area("Resume", st.session_state['saved_base'], height=400, label_visibility="collapsed")
    with c2:
        st.subheader("\U0001f4bc Job Description")
        st.caption("Paste the full JD here.")
        jd = st.text_area("JD", st.session_state['saved_jd'], height=400, label_visibility="collapsed")
    if st.button("\u2728 Generate Tailored Resume", type="primary", use_container_width=True):
        if base and jd and google_key:
            st.session_state['saved_base'] = base
            st.session_state['saved_jd'] = jd
            with st.spinner("Harvesting keywords, bridging stack, optimising for ATS..."):
                data = analyze_and_generate(google_key, base, jd)
                if "error" in data:
                    st.error(data['error'])
                else:
                    st.session_state['data'] = data
                    st.rerun()
        else:
            st.warning("Please provide API Key and paste a Job Description.")
else:
    data = st.session_state['data']
    # Top bar
    c1, c2, c3 = st.columns([1, 4, 1])
    with c2:
        st.markdown(f"## \U0001f3af Target: {data.get('target_company', 'Company')}")
    with c3:
        score = data.get('ats_score', 0)
        st.metric("ATS Match", f"{score}%")
    # Missing keywords alert
    missing = data.get('missing_keywords', '')
    if missing and str(missing).strip():
        st.warning(f"**Keywords still missing from resume:** {missing}")
    tab_edit, tab_export, tab_cover = st.tabs(["\U0001f4dd Editor", "\U0001f680 Export", "\u270d\ufe0f Cover Letter"])
    with tab_edit:
        with st.form("edit_form"):
            st.subheader("Candidate Details")
            c1, c2, c3 = st.columns(3)
            data['candidate_name'] = c1.text_input("Name", to_text_block(data.get('candidate_name')))
            data['candidate_title'] = c2.text_input("Title", to_text_block(data.get('candidate_title')))
            data['contact_info'] = c3.text_input("Contact", to_text_block(data.get('contact_info')))
            data['summary'] = st.text_area("Summary", to_text_block(data.get('summary')), height=120)
            st.subheader("Skills")
            skills = data.get('skills', {})
            new_skills = {}
            s_cols = st.columns(2)
            for i, (k, v) in enumerate(skills.items()):
                col = s_cols[i % 2]
                new_val = col.text_area(k, to_text_block(v), key=f"skill_{i}", height=80)
                new_skills[k] = new_val.replace('\n', ', ')
            data['skills'] = new_skills
            st.subheader("Experience")
            for i, role in enumerate(data.get('experience', [])):
                with st.expander(f"{role.get('role_title', 'Role')} @ {role.get('company', 'Company')}"):
                    c1, c2 = st.columns(2)
                    role['role_title'] = c1.text_input("Title", to_text_block(role.get('role_title')), key=f"jt_{i}")
                    role['company'] = c2.text_input("Company", to_text_block(role.get('company')), key=f"jc_{i}")
                    c3, c4 = st.columns(2)
                    role['dates'] = c3.text_input("Dates", to_text_block(role.get('dates')), key=f"jd_{i}")
                    role['location'] = c4.text_input("Location", to_text_block(role.get('location')), key=f"jl_{i}")
                    role['responsibilities'] = st.text_area("Responsibilities", to_text_block(role.get('responsibilities')), height=250, key=f"jr_{i}")
                    role['achievements'] = st.text_area("Achievements", to_text_block(role.get('achievements')), height=100, key=f"ja_{i}")
            st.subheader("Education")
            for i, edu in enumerate(data.get('education', [])):
                c1, c2 = st.columns(2)
                edu['degree'] = c1.text_input("Degree", to_text_block(edu.get('degree')), key=f"ed_{i}")
                edu['college'] = c2.text_input("Institution", to_text_block(edu.get('college')), key=f"ec_{i}")
            if st.form_submit_button("\U0001f4be Save Edits", type="primary"):
                st.session_state['data'] = data
                st.success("Saved!")
                st.rerun()
    with tab_export:
        st.subheader("\U0001f4e5 Download")
        c_name = data.get('candidate_name', 'Akhila_Reddy')
        default_company = data.get('target_company', 'Company')
        target_company = st.text_input("Company (for filename)", default_company)
        safe_name = re.sub(r'[^a-zA-Z0-9_-]', '_', c_name.strip().replace(' ', '_'))
        safe_company = re.sub(r'[^a-zA-Z0-9_-]', '_', target_company.strip())
        final_filename = f"{safe_name}_{safe_company}"
        c1, c2 = st.columns(2)
        doc_obj = create_doc(data)
        bio = io.BytesIO()
        doc_obj.save(bio)
        c1.download_button(
            label="\U0001f4c4 Word (.docx)",
            data=bio.getvalue(),
            file_name=f"{final_filename}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=True,
        )
        try:
            pdf_data = create_pdf(data)
            c2.download_button(
                label="\U0001f4d5 PDF",
                data=pdf_data,
                file_name=f"{final_filename}.pdf",
                mime="application/pdf",
                type="secondary",
                use_container_width=True,
            )
        except Exception as e:
            c2.error(f"PDF Error: {e}")
    with tab_cover:
        st.subheader("\u270d\ufe0f Cover Letter")
        st.info("Generates a human-sounding cover letter using your best matching war story for this JD.")
        if st.button("\u2728 Draft Cover Letter", type="primary"):
            if google_key and st.session_state['saved_jd']:
                with st.spinner("Picking war story, drafting narrative..."):
                    cl_text = generate_cover_letter(google_key, data, st.session_state['saved_jd'])
                    st.session_state['cover_letter'] = cl_text
            else:
                st.warning("Need API key and JD.")
        if st.session_state['cover_letter']:
            edited_cl = st.text_area("Preview (editable)", st.session_state['cover_letter'], height=400)
            st.session_state['cover_letter'] = edited_cl
            cl_doc = create_cover_letter_doc(st.session_state['cover_letter'], data)
            bio_cl = io.BytesIO()
            cl_doc.save(bio_cl)
            st.download_button(
                label="\U0001f4c4 Download Cover Letter (.docx)",
                data=bio_cl.getvalue(),
                file_name=f"Cover_Letter_{final_filename}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
            )
    st.divider()
    c3, c4 = st.columns(2)
    if c3.button("\u267b\ufe0f Re-Optimise", use_container_width=True):
        if st.session_state['saved_base'] and st.session_state['saved_jd']:
            with st.spinner("Re-tailoring..."):
                data = analyze_and_generate(google_key, st.session_state['saved_base'], st.session_state['saved_jd'])
                if "error" in data:
                    st.error(data['error'])
                else:
                    st.session_state['data'] = data
                    st.rerun()
    if c4.button("New Application (Keep Resume)", use_container_width=True):
        st.session_state['data'] = None
        st.session_state['saved_jd'] = ""
        st.session_state['cover_letter'] = None
        st.rerun()
